from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import pytesseract
from PIL import Image, ImageOps
import io
import re
from pdf2image import convert_from_bytes
import cv2
import numpy as np


def extraer_texto_pdf(file_bytes):
    # Intentar extracción por texto (PDF no escaneado)
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texto = "".join([page.extract_text() or "" for page in reader.pages])
        return texto
    except Exception as e:
        return ""


def extraer_texto_imagen(file_bytes):
    image = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(image, lang="spa")


# extraer_texto_word
def extraer_texto_word(file_bytes) -> str:
    """
    Extrae texto de un archivo Word (.docx) dado como bytes.
    """
    doc = Document(io.BytesIO(file_bytes))
    lineas = []
    # tratamiento si hay tablas
    for table in doc.tables:
        for row in table.rows:
            celdas = [c.text.strip() for c in row.cells]
            lineas.append(" | ".join(celdas))
    # tratamiento si hay parrafos normales
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto:
            lineas.append(texto)
    texto_completo = "\n".join(lineas)
    if "ANEXO" in texto:
        partes = texto.split("ANEXO")
        texto = "ANEXO".join(partes[-2:]) + "\n\n" + texto
    return texto_completo


def extraer_texto_excel(file_bytes):
    """Extrae texto de un archivo Excel (.xls, .xlsx) dado como bytes."""
    try:
        xls = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        contenido = ""
        for nombre, hoja in xls.items():
            contenido += f"--- Hoja: {nombre} ---\n"
            contenido += hoja.to_string(index=False)
        return contenido
    except Exception as e:
        return f"Error leyendo Excel: {e}"


# --------- funcionalidades OCR ---------
def auto_rotate(image: Image.Image, lang: str = "spa") -> Image.Image:
    """
    Detecta automáticamente la orientación del texto y devuelve la imagen corregida.
    """
    try:
        angles = [0, 90, 180, 270]
        scores = []

        for angle in angles:
            rotated = image.rotate(angle, expand=True)
            text = pytesseract.image_to_string(rotated, lang=lang)
            # mide cantidad de texto válido
            scores.append((len(text.strip()), angle))

        best_angle = max(scores, key=lambda x: x[0])[1]
        if best_angle != 0:
            return image.rotate(best_angle, expand=True)
        return image
    except Exception:
        return image


def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    img = auto_rotate(img)
    g = ImageOps.grayscale(img)
    if g.width < 1000:
        ratio = 1000 / float(g.width)
        g = g.resize((1000, int(g.height * ratio)), Image.LANCZOS)
    g = ImageOps.autocontrast(g)
    return g


def _ocr_image(img: Image.Image, lang: str = "spa", psm: int = 6) -> str:
    config = f"--psm 3 --oem 1"
    return pytesseract.image_to_string(img, lang=lang, config=config)


def extraer_texto_pdf_ocr(
    file_bytes: bytes, lang: str = "spa", dpi: int = 120, max_pages: int = 10
) -> str:
    """
    extrae texto mediante OCR página a página
    """
    try:
        pages = convert_from_bytes(
            file_bytes, dpi=dpi, first_page=1, last_page=max_pages
        )
        texto_total = []
        for i, page in enumerate(pages):
            if i >= max_pages:
                break
            proc = _preprocess_for_ocr(page)
            texto = _ocr_image(proc, lang=lang, psm=6).strip()
            texto_total.append(f"--- Página {i+1} ---\n{texto}")
            page.close()
            proc.close()
        texto_total = "\n\n".join(texto_total)
        # Limpieza básica
        texto_total = re.sub(r"\s+", " ", texto_total)
        texto_total = re.sub(r"[^\w\sáéíóúñÁÉÍÓÚÑ.,:;()\-/$%#°]", "", texto_total)
        return texto_total.strip()
    except Exception as e:
        return f"Error procesando el archivo, demasiadas paginas: {str(e)}"


def extraer_texto_imagen(file_bytes: bytes, lang: str = "spa") -> str:
    """
    OCR de una imagen.
    """
    image = Image.open(io.BytesIO(file_bytes))
    proc = _preprocess_for_ocr(image)
    texto = _ocr_image(proc, lang=lang, psm=6)
    image.close()
    proc.close()
    # Limpieza basica
    texto = re.sub(r"\s+", " ", texto)
    texto = re.sub(r"[^\w\sáéíóúñÁÉÍÓÚÑ.,:;()\-/$%#°]", "", texto)
    return texto.strip()


# --------- Router principal ---------
def convertir_a_txt(file_bytes: bytes, extension: str, lang_ocr: str = "spa") -> str:
    """
    Enruta por extensión. Para PDF: intenta texto embebido y cae a OCR si es corto.
    """
    ext = (extension or "").lower()
    if ext == ".pdf":
        texto_directo = extraer_texto_pdf(file_bytes)
        if len((texto_directo or "").strip()) < 1000:
            texto_ocr = extraer_texto_pdf_ocr(
                # limita la cantidad de páginas a procesar en 30 y dpi a 150
                file_bytes,
                lang=lang_ocr,
                dpi=120,
                max_pages=10,
            )
            # Conserva el más largo y con más contenido
            if len(texto_ocr) > len(texto_directo or ""):
                return texto_ocr
        return texto_directo or ""
    elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"):
        return extraer_texto_imagen(file_bytes, lang=lang_ocr)
    elif ext in (".docx",):
        return extraer_texto_word(file_bytes)
    elif ext in (".xls", ".xlsx"):
        return extraer_texto_excel(file_bytes)
    else:
        return extraer_texto_imagen(file_bytes, lang=lang_ocr)
